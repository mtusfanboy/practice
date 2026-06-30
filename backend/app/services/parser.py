"""Извлечение текста из документов (BE-04).

Поддерживаются форматы PDF (через ``pdfplumber``) и DOCX (через
``python-docx``). Текст извлекается постранично, что позволяет
сохранять номер страницы для каждого фрагмента (BE-07).
"""

import io
from dataclasses import dataclass

import pdfplumber
from docx import Document as DocxDocument

from app.core.logging import get_logger

logger = get_logger(__name__)


class DocumentParsingError(Exception):
    """Ошибка извлечения текста из документа."""


@dataclass(frozen=True)
class ParsedPage:
    """Извлечённая страница документа.

    :param page_number: номер страницы, начиная с 1.
    :param text: текстовое содержимое страницы.
    """

    page_number: int
    text: str


def parse_document(content: bytes, file_name: str) -> list[ParsedPage]:
    """Извлечь текст из документа в зависимости от его расширения.

    :param content: бинарное содержимое файла.
    :param file_name: имя файла (используется для определения формата).
    :return: список страниц с извлечённым текстом.
    :raises DocumentParsingError: если формат не поддерживается или
        извлечение текста завершилось ошибкой.
    """
    extension = file_name.rsplit(".", 1)[-1].lower() if "." in file_name else ""

    if extension == "pdf":
        return _parse_pdf(content)
    if extension == "docx":
        return _parse_docx(content)

    raise DocumentParsingError(
        f"Неподдерживаемый формат файла: '.{extension}'. "
        "Поддерживаются только PDF и DOCX."
    )


def _parse_pdf(content: bytes) -> list[ParsedPage]:
    """Извлечь текст из PDF постранично с помощью ``pdfplumber``.

    :param content: бинарное содержимое PDF-файла.
    :return: список непустых страниц с текстом.
    :raises DocumentParsingError: если файл повреждён или не читается.
    """
    pages: list[ParsedPage] = []
    try:
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            for index, page in enumerate(pdf.pages, start=1):
                text = page.extract_text() or ""
                text = text.strip()
                if text:
                    pages.append(ParsedPage(page_number=index, text=text))
    except DocumentParsingError:
        raise
    except Exception as exc:  # noqa: BLE001 — оборачиваем любую ошибку библиотеки
        raise DocumentParsingError(
            f"Не удалось извлечь текст из PDF: {exc}"
        ) from exc

    if not pages:
        raise DocumentParsingError(
            "PDF не содержит извлекаемого текста (возможно, это скан-изображение)."
        )
    return pages


def _parse_docx(content: bytes) -> list[ParsedPage]:
    """Извлечь текст из DOCX-файла с помощью ``python-docx``.

    DOCX не имеет явного деления на страницы, поэтому весь текст
    относится к странице номер 1.

    :param content: бинарное содержимое DOCX-файла.
    :return: список из одной страницы с собранным текстом.
    :raises DocumentParsingError: если файл повреждён или не читается.
    """
    try:
        document = DocxDocument(io.BytesIO(content))
    except Exception as exc:  # noqa: BLE001
        raise DocumentParsingError(
            f"Не удалось открыть DOCX-файл: {exc}"
        ) from exc

    parts: list[str] = [
        paragraph.text.strip()
        for paragraph in document.paragraphs
        if paragraph.text and paragraph.text.strip()
    ]

    # Также извлекаем текст из таблиц.
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    text = "\n".join(parts).strip()
    if not text:
        raise DocumentParsingError("DOCX-файл не содержит текста.")

    return [ParsedPage(page_number=1, text=text)]
