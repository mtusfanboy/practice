"""Генератор набора тестовых документов (QA-03).

Создаёт в каталоге ``tests/fixtures`` репрезентативный набор файлов для
тестирования парсинга и индексации:

* ``sample_lecture.pdf`` — корректный многостраничный PDF на русском языке;
* ``sample_lecture.docx`` — корректный DOCX с абзацами и таблицей;
* ``custom_font.pdf`` — PDF с нестандартным (моноширинным) шрифтом;
* ``empty.pdf`` / ``empty.docx`` — валидные файлы без извлекаемого текста;
* ``corrupted.pdf`` / ``corrupted.docx`` — файлы с битым форматированием;
* ``not_allowed.txt`` — файл неподдерживаемого формата.

Скрипт идемпотентен: повторный запуск перезаписывает файлы. Запуск::

    python tests/fixtures/generate_fixtures.py
"""

import os
from pathlib import Path

import reportlab
from docx import Document as DocxDocument
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas

FIXTURES_DIR = Path(__file__).resolve().parent

# Кандидаты шрифтов с поддержкой кириллицы для разных ОС. Базовые шрифты PDF
# (Helvetica/Courier/Times) не содержат кириллических глифов, поэтому для
# русского текста необходим встроенный TrueType-шрифт.
_CYRILLIC_FONT_CANDIDATES = [
    "/System/Library/Fonts/Supplemental/Arial.ttf",  # macOS
    "/Library/Fonts/Arial.ttf",  # macOS (пользовательские)
    "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",  # Debian/Ubuntu
    "/usr/share/fonts/dejavu/DejaVuSans.ttf",  # Fedora/RHEL
    "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf",
    "/usr/share/fonts/TTF/DejaVuSans.ttf",  # Arch
]

# «Нестандартный» встроенный шрифт — поставляется вместе с reportlab,
# поэтому доступен на любой платформе (используется для custom_font.pdf).
_VERA_BOLD = os.path.join(os.path.dirname(reportlab.__file__), "fonts", "VeraBd.ttf")

CYRILLIC_FONT = "CyrillicFont"
CUSTOM_FONT = "CustomEmbeddedFont"


def _register_fonts() -> str:
    """Зарегистрировать шрифты для генерации PDF.

    :return: имя зарегистрированного кириллического шрифта.
    :raises RuntimeError: если ни один кириллический шрифт не найден.
    """
    cyrillic_path = next(
        (path for path in _CYRILLIC_FONT_CANDIDATES if os.path.exists(path)), None
    )
    if cyrillic_path is None:
        raise RuntimeError(
            "Не найден TrueType-шрифт с поддержкой кириллицы. "
            "Установите шрифты DejaVu (fonts-dejavu-core) или Arial."
        )
    pdfmetrics.registerFont(TTFont(CYRILLIC_FONT, cyrillic_path))
    pdfmetrics.registerFont(TTFont(CUSTOM_FONT, _VERA_BOLD))
    return CYRILLIC_FONT

# Содержательный русскоязычный текст для проверки русского анализатора.
LECTURE_PAGES = [
    (
        "Лекция 1. Введение в базы данных. "
        "Реляционная модель данных была предложена Эдгаром Коддом в 1970 году. "
        "Основными понятиями являются отношение, кортеж и атрибут. "
        "Язык структурированных запросов SQL используется для управления данными. "
        "Нормализация позволяет устранить избыточность и аномалии обновления."
    ),
    (
        "Лекция 2. Индексы и оптимизация запросов. "
        "Индекс ускоряет поиск строк по значениям столбцов. "
        "B-дерево является наиболее распространённой структурой индекса. "
        "Полнотекстовый поиск опирается на инвертированный индекс. "
        "Кеширование часто запрашиваемых данных снижает нагрузку на систему."
    ),
]


def _write_pdf(path: Path, pages: list[str], font_name: str = CYRILLIC_FONT) -> None:
    """Создать PDF, размещая каждый элемент списка на отдельной странице.

    :param path: путь к создаваемому файлу.
    :param pages: список текстов страниц.
    :param font_name: имя зарегистрированного шрифта.
    """
    pdf = canvas.Canvas(str(path), pagesize=A4)
    width, height = A4
    for text in pages:
        pdf.setFont(font_name, 12)
        text_object = pdf.beginText(50, height - 60)
        # Разбиваем длинный текст на строки фиксированной ширины.
        words = text.split(" ")
        line = ""
        for word in words:
            if len(line) + len(word) + 1 > 70:
                text_object.textLine(line)
                line = word
            else:
                line = f"{line} {word}".strip()
        if line:
            text_object.textLine(line)
        pdf.drawText(text_object)
        pdf.showPage()
    pdf.save()


def _write_docx(path: Path) -> None:
    """Создать корректный DOCX с заголовком, абзацами и таблицей."""
    document = DocxDocument()
    document.add_heading("Конспект лекций по базам данных", level=1)
    for text in LECTURE_PAGES:
        document.add_paragraph(text)

    table = document.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Понятие"
    table.rows[0].cells[1].text = "Определение"
    table.rows[1].cells[0].text = "Индекс"
    table.rows[1].cells[1].text = "Структура для ускорения поиска"
    document.save(str(path))


def _write_empty_docx(path: Path) -> None:
    """Создать валидный DOCX без текстового содержимого."""
    document = DocxDocument()
    document.save(str(path))


def _write_empty_pdf(path: Path) -> None:
    """Создать валидный PDF с пустой страницей (без извлекаемого текста)."""
    pdf = canvas.Canvas(str(path), pagesize=A4)
    pdf.showPage()
    pdf.save()


def generate_all() -> dict[str, Path]:
    """Сгенерировать весь набор фикстур.

    :return: словарь ``{имя_файла: путь}`` для созданных файлов.
    """
    FIXTURES_DIR.mkdir(parents=True, exist_ok=True)
    _register_fonts()

    paths: dict[str, Path] = {}

    sample_pdf = FIXTURES_DIR / "sample_lecture.pdf"
    _write_pdf(sample_pdf, LECTURE_PAGES)
    paths["sample_lecture.pdf"] = sample_pdf

    # PDF с нестандартным встроенным шрифтом (латиница) — проверка того,
    # что парсер справляется с произвольными встроенными шрифтами (QA-03).
    custom_font_pdf = FIXTURES_DIR / "custom_font.pdf"
    _write_pdf(
        custom_font_pdf,
        ["Custom embedded font sample text 1234567890 ABCDEFG."],
        font_name=CUSTOM_FONT,
    )
    paths["custom_font.pdf"] = custom_font_pdf

    sample_docx = FIXTURES_DIR / "sample_lecture.docx"
    _write_docx(sample_docx)
    paths["sample_lecture.docx"] = sample_docx

    empty_pdf = FIXTURES_DIR / "empty.pdf"
    _write_empty_pdf(empty_pdf)
    paths["empty.pdf"] = empty_pdf

    empty_docx = FIXTURES_DIR / "empty.docx"
    _write_empty_docx(empty_docx)
    paths["empty.docx"] = empty_docx

    corrupted_pdf = FIXTURES_DIR / "corrupted.pdf"
    corrupted_pdf.write_bytes(b"%PDF-1.4 this is not a valid pdf body \x00\x01\x02")
    paths["corrupted.pdf"] = corrupted_pdf

    corrupted_docx = FIXTURES_DIR / "corrupted.docx"
    corrupted_docx.write_bytes(b"PK\x03\x04 broken docx content not a real zip")
    paths["corrupted.docx"] = corrupted_docx

    not_allowed = FIXTURES_DIR / "not_allowed.txt"
    not_allowed.write_text("Текстовый файл неподдерживаемого формата.", encoding="utf-8")
    paths["not_allowed.txt"] = not_allowed

    return paths


if __name__ == "__main__":
    created = generate_all()
    print(f"Создано {len(created)} тестовых файлов в {FIXTURES_DIR}:")
    for name in sorted(created):
        print(f"  - {name}")
